"""Service for extracting text from document files."""

import io
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader


class DocumentLoaderService:
    """Extracts plain text from uploaded documents (PDF, DOCX, Markdown, text)."""

    # Minimum average chars per page to consider PDF text extraction successful.
    # Below this, OCR fallback is attempted.
    _PDF_OCR_THRESHOLD = 20

    @staticmethod
    def extract_text(filename: str, content: bytes) -> str:
        """Extract text from file content based on file extension.

        Args:
            filename: Original filename (used to determine format).
            content: Raw file bytes.

        Returns:
            Extracted text. PDFs use pypdf (with OCR fallback if needed);
            DOCX uses python-docx; other files are decoded as UTF-8.
        """
        suffix = Path(filename).suffix.lower()
        if suffix == ".pdf":
            return DocumentLoaderService._extract_pdf_text(content)
        if suffix in (".docx", ".doc"):
            return DocumentLoaderService._extract_docx_text(content)
        return content.decode("utf-8", errors="ignore")

    @staticmethod
    def _extract_pdf_text(content: bytes) -> str:
        """Extract text from PDF bytes using pypdf, with OCR fallback for scanned PDFs.

        Args:
            content: Raw PDF file bytes.

        Returns:
            Concatenated text from all pages, stripped of leading/trailing whitespace.
        """
        reader = PdfReader(io.BytesIO(content))
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n".join(pages).strip()

        # Fallback to OCR if text extraction yielded little content (e.g. scanned PDF)
        if pages and DocumentLoaderService._should_use_pdf_ocr(pages):
            ocr_text = DocumentLoaderService._extract_pdf_ocr(content)
            if ocr_text.strip():
                return ocr_text

        return text

    @staticmethod
    def _should_use_pdf_ocr(pages: list[str]) -> bool:
        """Return True if average chars per page is below threshold (likely scanned)."""
        total = sum(len(p.strip()) for p in pages)
        return len(pages) > 0 and (total / len(pages)) < DocumentLoaderService._PDF_OCR_THRESHOLD

    @staticmethod
    def _extract_pdf_ocr(content: bytes) -> str:
        """Extract text from PDF using OCR (pdf2image + pytesseract).

        Args:
            content: Raw PDF file bytes.

        Returns:
            OCR-extracted text from all pages, or empty string if OCR fails.
        """
        try:
            from pdf2image import convert_from_bytes
            import pytesseract
        except ImportError:
            return ""

        try:
            images = convert_from_bytes(content)
            texts = [pytesseract.image_to_string(img) for img in images]
            return "\n".join(texts).strip()
        except Exception:
            return ""

    @staticmethod
    def _extract_docx_text(content: bytes) -> str:
        """Extract text from DOCX bytes using python-docx.

        Args:
            content: Raw DOCX file bytes.

        Returns:
            Extracted text from paragraphs and tables.
        """
        doc = DocxDocument(io.BytesIO(content))
        parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)

        return "\n".join(parts).strip()
